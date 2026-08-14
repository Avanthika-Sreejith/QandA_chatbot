from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "Document_QA_Chatbot_Updated_Project_Report.docx"
FLOW = ROOT / "work" / "document_qa_updated_flow.png"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
PALE = "E8EEF5"
LIGHT = "F2F4F7"
GREEN = "EAF4ED"
GOLD = "7A5A00"
RED = "9B1C1C"
PAGE_W = Inches(6.5)


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(inches * 1440)))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    ind = tbl_pr.first_child_found_in("w:tblInd")
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def font(run, size=11, color="000000", bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_text(doc, text="", size=11, color="000000", bold=False, italic=False, after=6, before=0, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    font(r, size, color, bold, italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        font(r, 11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        font(r, 11)


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_table(doc, headers, rows, widths, small=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, headers):
        set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        font(r, 9.3 if small else 10, NAVY, True)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(text))
            font(r, 8.8 if small else 9.5)
    return table


def page_field(paragraph):
    run = paragraph.add_run("Page ")
    font(run, 9, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Document Q&A Chatbot | Technical Project Report")
    font(run, 9, MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_field(footer)


def draw_flow():
    img = Image.new("RGB", (1800, 970), "white")
    d = ImageDraw.Draw(img)
    try:
        h1 = ImageFont.truetype("arialbd.ttf", 34)
        h2 = ImageFont.truetype("arial.ttf", 24)
        sm = ImageFont.truetype("arial.ttf", 19)
    except OSError:
        h1 = h2 = sm = ImageFont.load_default()
    navy = "#0B2545"; blue = "#2E74B5"; pale = "#E8EEF5"; green = "#EAF4ED"; gray = "#536171"
    d.text((65, 38), "Updated document Q&A processing flow", font=h1, fill=navy)
    nodes = [
        (70, 180, 300, 110, "Upload files\nor ZIP folder", pale),
        (445, 180, 300, 110, "Safe extraction\n+ parsing", pale),
        (820, 180, 300, 110, "Section-aware\nsegments", pale),
        (1195, 180, 300, 110, "Semantic chunks\n+ metadata", pale),
        (620, 490, 320, 115, "Jina API dense vectors\n+ FastEmbed BM25", green),
        (1110, 490, 320, 115, "Qdrant Cloud\ndense + sparse + payload", green),
        (190, 765, 300, 110, "User question", pale),
        (645, 745, 340, 150, "Hybrid retrieval\nRRF + source selection\n+ sequence expansion", pale),
        (1200, 765, 350, 110, "Groq answer + inline\ncitations + chunk viewer", pale),
    ]
    for x,y,w,h,text,fill in nodes:
        d.rounded_rectangle((x,y,x+w,y+h), 18, fill=fill, outline=blue, width=3)
        lines=text.split("\n")
        line_h=28
        ty=y+(h-len(lines)*line_h)//2-4
        for line in lines:
            bb=d.textbbox((0,0), line, font=h2)
            d.text((x+(w-(bb[2]-bb[0]))//2, ty), line, font=h2, fill=navy)
            ty += line_h
    def arrow(a,b):
        d.line((a[0],a[1],b[0],b[1]), fill=gray, width=5)
        d.polygon([(b[0],b[1]),(b[0]-15,b[1]-8),(b[0]-15,b[1]+8)], fill=gray)
    arrow((370,235),(440,235)); arrow((745,235),(815,235)); arrow((1120,235),(1190,235))
    arrow((1345,290),(825,480)); arrow((940,547),(1100,547))
    arrow((490,820),(635,820)); arrow((985,820),(1190,820))
    d.text((70, 370), "Files are retained under their document chat, so another chat is never searched.", font=sm, fill=gray)
    d.text((70, 660), "Question embedding uses the same Jina API. Qdrant returns evidence, then Groq writes only from those chunks.", font=sm, fill=gray)
    FLOW.parent.mkdir(parents=True, exist_ok=True)
    img.save(FLOW)


def cover(doc):
    add_text(doc, "PROJECT REPORT", 11, GOLD, True, after=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Document Q&A Chatbot", 30, NAVY, True, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Updated Architecture, Implementation and Validation Report", 15, DARK_BLUE, False, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "A chat-scoped Retrieval-Augmented Generation (RAG) application for answering questions from uploaded documents.", 11, MUTED, False, after=36, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_table(doc, ["Report item", "Current implementation"], [
        ("Application", "Document Q&A Chatbot"),
        ("User interface", "Streamlit"),
        ("Deployment", "Render with Docker"),
        ("Vector database", "Qdrant Cloud"),
        ("Chat metadata", "Supabase"),
        ("Report date", date.today().strftime("%d %B %Y")),
    ], [1.75, 4.75])
    add_text(doc, "Purpose", 11, NAVY, True, before=28, after=3)
    add_text(doc, "This report explains the updated system in plain language, documents the changes made during development, and records the design choices that improve accuracy, source isolation, citations, and deployment reliability.", 11, "000000", after=0)
    doc.add_page_break()


def main():
    draw_flow()
    doc = Document()
    setup(doc)
    cover(doc)

    add_heading(doc, "1. Executive Summary")
    add_text(doc, "The Document Q&A Chatbot lets a user create separate document chats, upload files or a ZIP archive of a folder, index the content, and ask questions that are answered only from the files in the active chat. It is a Retrieval-Augmented Generation (RAG) system: the language model does not answer directly from general memory; it receives relevant retrieved passages and produces an answer with citations.")
    add_text(doc, "The current version replaces the earlier local, model-heavy workflow with API-based dense embeddings through Jina, hybrid retrieval in Qdrant Cloud, and Groq for answer generation. It also adds semantic chunking, automatic source selection, section-aware parsing, precise citations, ZIP folder upload, persistent document chats, and a safe OCR switch for low-memory deployment.")

    add_heading(doc, "2. Problem Statement and Objectives")
    add_text(doc, "A typical document collection may contain PDFs, Word files, spreadsheets and text files. Searching manually is slow, while a generic chatbot can produce an answer without showing where it came from. The system addresses this by turning uploaded content into searchable chunks while keeping each document chat isolated.")
    add_table(doc, ["Objective", "How the current system addresses it"], [
        ("Answer from the right documents", "Every Qdrant search is filtered by document_chat_id. The active chat is the only search scope."),
        ("Preserve documents for later", "Files remain indexed in their chat; users can switch chats, rename them, inspect indexed files, or delete a chat."),
        ("Avoid unrelated file sources", "Two-stage source selection selects the strongest file(s), then re-searches only inside them."),
        ("Keep related content together", "Semantic chunking uses sentence embeddings and preserves definition lists, tables and procedures."),
        ("Make answers traceable", "The response includes inline [n] citations, a Sources list, and an expandable view of exact retrieved chunks."),
        ("Support deployment", "Qdrant is cloud-hosted; Render runs the Streamlit app in Docker. OCR is configurable to fit the memory plan."),
    ], [2.2, 4.3])

    add_heading(doc, "3. Current System Architecture")
    add_text(doc, "The following diagram shows the current end-to-end flow. All vector storage is remote in Qdrant Cloud; Supabase stores the chat and question-history data. The hosting machine does not need the developer laptop's GPU.")
    doc.add_picture(str(FLOW), width=PAGE_W)
    p = add_text(doc, "Figure 1. Updated ingestion and question-answering flow.", 9, MUTED, italic=True, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "4. Technology Stack")
    add_table(doc, ["Layer", "Technology", "Responsibility"], [
        ("Web application", "Streamlit", "User interface for chats, uploads, indexing progress, answers and retrieved-chunk inspection."),
        ("Document parsers", "PyMuPDF, python-docx, openpyxl", "Extract PDF, DOCX, XLSX/XLSM and TXT content plus citation metadata."),
        ("Dense embedding", "Jina Embeddings v3 API", "Produces 1024-dimensional semantic vectors for chunks and questions."),
        ("Sparse embedding", "FastEmbed BM25 (Qdrant/bm25)", "Provides keyword-aware sparse vectors for hybrid retrieval."),
        ("Vector database", "Qdrant Cloud", "Stores dense/sparse vectors and filterable payload metadata."),
        ("Answer generator", "Groq API (Llama 3.3 70B)", "Writes grounded answers from retrieved chunks and returns citation references."),
        ("Chat persistence", "Supabase", "Stores chat names, document-chat records and question/answer history."),
        ("Deployment", "Render + Docker", "Runs the hosted Streamlit service and installs required system libraries."),
    ], [1.35, 1.85, 3.3], small=True)

    add_heading(doc, "5. What Changed During Development")
    add_text(doc, "The project was improved iteratively after observing real retrieval, usability and hosting issues. The main changes are summarised below.")
    add_table(doc, ["Earlier state", "Current state", "Reason / benefit"], [
        ("Local Qdrant Docker", "Qdrant Cloud", "Removes dependency on a local Docker database and supports hosted deployment."),
        ("Fixed-size character chunks", "Embedding-driven semantic chunks", "Reduces topic cuts and preserves related explanations together."),
        ("Single retrieval signal", "Dense Jina + sparse BM25 hybrid retrieval with RRF", "Balances semantic meaning with exact terms, headings and formulas."),
        ("All uploaded files could appear as sources", "Automatic best-source selection", "An unrelated file is excluded unless its evidence is close to the strongest source."),
        ("Simple file-level source display", "Inline citations plus exact chunk viewer", "Makes the answer auditable at page, section, sheet/table-row and chunk level."),
        ("Files only", "Files plus safe ZIP folder upload", "Allows a folder and subfolders to be uploaded in a hosting-compatible way."),
        ("OCR loaded by default", "OCR disabled by default on 512 MB Render", "Prevents RapidOCR/ONNX from exhausting service memory; can be enabled on a larger plan."),
        ("Generic chat labels", "Automatic filename-based title, manual rename, delete and indexed-files view", "Makes separate document chats understandable and manageable."),
    ], [1.7, 2.05, 2.75], small=True)

    add_heading(doc, "6. Ingestion Workflow")
    add_numbered(doc, [
        "The user creates or selects a document chat. The active chat ID becomes the mandatory Qdrant filter for indexing and search.",
        "The user uploads one or more supported documents, or uploads a ZIP archive containing a folder. ZIP paths are checked to reject path traversal, and file-count and uncompressed-size limits are enforced.",
        "The parser extracts source segments and metadata. A segment can be prose, a heading, a table or an image-related unit.",
        "Semantic chunking turns prose into topic-coherent retrieval units. Tables and image units are retained as their own source units rather than split into arbitrary characters.",
        "Each chunk receives a Jina dense vector and a FastEmbed BM25 sparse vector. The text and metadata are stored together as one Qdrant point.",
        "The app reports indexed file, segment and chunk counts. The upload widget resets after success so already-indexed files are not mistakenly shown as pending again.",
    ])

    add_heading(doc, "6.1 Supported Uploads and Safe ZIP Folders", 2)
    add_table(doc, ["Input", "Handling"], [
        ("PDF", "Native text, headings and tables are extracted. OCR is optional and disabled by default in the 512 MB Render configuration."),
        ("DOCX", "Paragraphs, headings and tables are read. Image-only DOCX content can return zero text segments."),
        ("XLSX / XLSM", "Worksheet data is converted into header-aware row groups. Formula expressions and available cached values are retained when present."),
        ("TXT", "Text is converted into source segments and semantically chunked."),
        ("ZIP", "Supported documents in nested folders are extracted safely; unsupported files are ignored. Limits: 1,000 members and 500 MB extracted supported content."),
        ("PNG / JPG direct upload", "Rejected by the file type filter in the current version."),
    ], [1.5, 5.0])

    add_heading(doc, "6.2 Table and Formula Handling", 2)
    add_text(doc, "Tables are not treated as ordinary paragraphs. The parser keeps table metadata (for example, page, table number and row range) and represents rows with their headers so that a retrieved result remains understandable. For spreadsheets, formula cells are read from a formula workbook and a values workbook, allowing the system to retain both a formula such as EV / AC and a cached calculated result if it exists in the file.")
    add_text(doc, "Formula images, scanned tables and diagrams require OCR or a vision/mathematical OCR service. This is intentionally not guaranteed on the 512 MB hosted plan because RapidOCR/ONNX models increase memory use significantly. A larger instance or a dedicated OCR/vision service is the recommended production option for visual formula extraction.", color=GOLD)

    add_heading(doc, "7. Semantic Chunking Design")
    add_text(doc, "Semantic chunking is different from fixed character splitting. The parser first separates useful sentences. Jina embeds those sentences in batches. The algorithm compares each next sentence to the mean vector of the current chunk and starts a new chunk when the semantic similarity falls below the configured threshold. The default threshold is 0.72, with a minimum of two sentences and a 3,500-character safety maximum.")
    add_table(doc, ["Rule", "Why it exists"], [
        ("Topic-shift boundary", "A similarity drop between the accumulated meaning and the next sentence suggests a new subtopic."),
        ("Numbered-item boundary", "A numbered step always starts a new chunk so a procedure is not mixed with preceding definitions."),
        ("Glossary-list join", "Consecutive lines such as Normal time, Crash time and Crash cost stay together so a key-concepts question retrieves the whole list."),
        ("Heading attachment", "Standalone headings are not indexed by themselves; they label the following section instead of being returned as an answer item."),
        ("Table/image unit", "Structured items remain intact rather than being cut through a row or visual unit."),
        ("Noise removal", "Page numbers, citation-only fragments, flow-chart labels and other extraction noise are skipped."),
    ], [1.8, 4.7])

    add_heading(doc, "8. Qdrant Storage Model")
    add_text(doc, "Qdrant stores one point per chunk. Each point has two vector representations and a payload containing the original text and citation metadata. Point IDs are deterministic UUIDs based on chat, file, source segment and chunk index, avoiding accidental duplicate records during repeat indexing.")
    add_text(doc, "Illustrative Qdrant point payload", 11, NAVY, True, before=4, after=3)
    sample = '''{
  "id": "deterministic UUID",
  "vector": {"dense": "1024 Jina values", "sparse": "BM25 indices + values"},
  "payload": {
    "document_chat_id": "chat UUID", "file_path": "Project_Management.pdf",
    "text": "Normal time: ... Crash time: ...", "page": 2,
    "section": "Crashing Project Activities", "source_segment": 12,
    "chunk_index": 0, "chunk_count": 1, "chunking_method": "semantic",
    "table": null, "table_row_start": null, "sheet": null
  }
}'''
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(.18)
    r = p.add_run(sample)
    font(r, 8.6, NAVY)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    add_text(doc, "Payload indexes are created for document_chat_id, file_path and source_segment. These indexes make chat isolation, file selection and neighbouring-procedure retrieval efficient.")

    add_heading(doc, "9. Question Answering and Citation Workflow")
    add_numbered(doc, [
        "The question is embedded using the Jina API and converted to a BM25 sparse vector using FastEmbed.",
        "Qdrant performs dense and sparse searches, always filtered to the active document_chat_id.",
        "Reciprocal Rank Fusion (RRF) combines both rankings. A chunk receives a contribution of 1 / (60 + rank) from each retrieval method where it appears.",
        "The system compares the best score from each file. Only files meeting the evidence threshold and at least 60% of the strongest file score are kept; it then re-searches only those file(s).",
        "For questions asking for steps, process, procedure or sequence, the system retrieves neighbouring chunks in the same source segment, in original chunk order. This prevents Step 1 from being omitted merely because Step 2 ranked higher.",
        "The selected chunks are sent to Groq with instructions to answer only from supplied evidence and include bracket citations such as [1].",
        "The interface renders the answer, Sources list and a View retrieved chunks panel showing the exact text, relevance score, file, page, section and chunk position.",
    ])

    add_heading(doc, "9.1 Why a citation may reference multiple chunks")
    add_text(doc, "A multi-point answer can correctly cite more than one chunk. For example, a list of definitions may be in one chunk while a rule about the critical path is in another. However, the citation should appear immediately after the statement it supports. The retrieved-chunk panel makes this check transparent. When parser or chunking logic changes, old points must be deleted and the file re-indexed; Qdrant does not automatically rewrite previously stored chunks.")

    add_heading(doc, "10. Chat Management and Persistence")
    add_table(doc, ["Feature", "Behaviour"], [
        ("Current chat scope", "Queries use only the active chat's document_chat_id filter; documents from another chat are excluded."),
        ("Automatic name", "A new chat can be named from the first uploaded file or ZIP folder name."),
        ("Manual rename", "The user can rename the document chat; the UI and stored record should refresh to show the new name."),
        ("Indexed files view", "The app lists files in the active chat so the user knows whether a document is already indexed."),
        ("Question history", "Supabase persists asked questions and answers; users can clear only the chat's displayed question history."),
        ("Delete chat", "Deletion removes the chat's Qdrant points using document_chat_id and removes corresponding Supabase data through the configured relationship/cascade."),
    ], [1.65, 4.85])

    add_heading(doc, "11. Deployment and Resource Design")
    add_text(doc, "The hosted application runs in Docker on Render. Qdrant Cloud, Jina, Groq and Supabase are accessed with API credentials stored as environment variables. This means vector database storage and answer/embedding computation are external services; the hosting service primarily runs parsing, chunk orchestration, sparse embeddings and the Streamlit user interface.")
    add_table(doc, ["Configuration", "Purpose"], [
        ("QDRANT_URL / QDRANT_API_KEY", "Connect to the Qdrant Cloud collection."),
        ("JINA_API_KEY", "Call the Jina OpenAI-compatible embedding API."),
        ("GROQ_API_KEY", "Generate grounded final answers."),
        ("SUPABASE_URL / SUPABASE_KEY", "Persist chats and question history."),
        ("EMBEDDING_DEVICE=cpu", "Appropriate for the hosted service because dense embeddings are API-based."),
        ("ENABLE_PDF_OCR=false", "Default on the 512 MB Render plan to avoid OCR/ONNX memory crashes."),
    ], [2.2, 4.3])
    add_text(doc, "GPU clarification: the developer laptop GPU can accelerate locally installed models, but it is not available to a hosted Render instance. In the current architecture, dense embeddings are already calculated by the Jina API, so the hosting CPU is not the main dense-embedding bottleneck. A hosted GPU requires a provider plan that explicitly supplies one.", color=GOLD)

    add_heading(doc, "12. Validation and Edge-Case Test Plan")
    add_text(doc, "The following cases reflect the edge cases tested or required during development. The report records expected observable behaviour rather than claiming a test passed without evidence.")
    add_table(doc, ["Scenario", "Expected behaviour / evidence"], [
        ("Empty PDF/DOCX/TXT", "Graceful message or zero-segment result; application must not crash."),
        ("Scanned or image-only PDF", "Clear notice when text/OCR is unavailable. On 512 MB Render, OCR remains disabled by design."),
        ("Large file (100 MB+)", "File-size/timeout feedback rather than silent failure; hosted plan limits must be observed."),
        ("Special characters in filename", "Parser preserves and handles Unicode/special characters safely."),
        ("Corrupt PDF/DOCX/XLSX", "Per-file error is shown without crashing the whole indexing session."),
        ("PNG/JPG direct upload", "Rejected by uploader type filter."),
        ("ZIP upload", "Safe extraction, supported-file filtering, nested-folder support, member and size limits."),
        ("10+ files", "Progress updates; valid files index without an unhandled timeout."),
        ("Same file twice in one chat", "Deterministic IDs make repeat-index behaviour predictable; duplicate-policy should be tested before production."),
        ("Blank / special-character query", "Prompted or safely handled; no unsafe query construction."),
        ("No indexed files", "Returns an empty, user-friendly result."),
        ("Question across document chats", "Only the active document_chat_id is searched."),
        ("Key concepts in crashing activities", "Definition lines should remain together and citations should identify the chunk containing those definitions."),
        ("Steps in a procedure", "Neighbouring chunks in the same segment are provided in source order so no leading step is lost."),
        ("Multi-file chat", "Unrelated file is suppressed unless it reaches the automatic source-selection threshold."),
    ], [2.35, 4.15], small=True)

    add_heading(doc, "13. Known Constraints and Recommended Next Steps")
    add_table(doc, ["Constraint", "Practical recommendation"], [
        ("Visual formulas and scanned table images", "Use a larger Render instance with ENABLE_PDF_OCR=true, or use a dedicated OCR/vision/math OCR API. Keep text-based tables on the current pipeline."),
        ("Changes to parser/chunker", "Delete and re-index affected documents. Existing Qdrant points contain the old extracted text and metadata."),
        ("Retrieval quality", "Continue testing difficult headings, lists, formulas and multi-file questions; tune threshold values only using representative evaluation questions."),
        ("Duplicate uploads", "Add a visible duplicate warning based on file hash if strict de-duplication becomes a requirement."),
        ("API availability/cost", "Monitor Jina and Groq API keys, rate limits and usage. Add user-facing errors for temporary API failures."),
        ("Citation validation", "Retain the retrieved-chunk viewer and spot-check that each [n] supports the sentence preceding it."),
    ], [2.05, 4.45])

    add_heading(doc, "14. Code Map for Understanding the Project")
    add_table(doc, ["File / module", "What to read there"], [
        ("app/streamlit_app.py", "UI flow, uploads, ZIP extraction, chat controls, progress feedback and answer rendering."),
        ("app/config.py", "All key thresholds, model names, vector configuration and the ENABLE_PDF_OCR switch."),
        ("app/parsers/documents.py", "PDF, DOCX, Excel and text extraction; headings, tables, formulas and optional OCR."),
        ("app/ingestion/chunking.py", "Sentence splitting, semantic similarity boundary logic, glossary/procedure safeguards and metadata propagation."),
        ("app/embeddings.py", "Jina dense API calls and FastEmbed BM25 vector creation."),
        ("app/ingestion/indexing.py", "Batch upsert of chunk text, vectors and deterministic point IDs to Qdrant."),
        ("app/retrieval/collection.py", "Qdrant collection schema and payload indexes."),
        ("app/retrieval/search.py", "Dense+sparse search, RRF fusion, source selection and sequence-context expansion."),
        ("app/generation/generator.py", "Grounded answer prompt and citation extraction."),
        ("app/database.py", "Supabase operations for chats and question history."),
        ("Dockerfile / render.yaml", "Hosted runtime, system libraries and deployment environment variables."),
    ], [2.1, 4.4])

    add_heading(doc, "15. Conclusion")
    add_text(doc, "The current Document Q&A Chatbot is no longer only a basic file-search prototype. It is a chat-scoped hybrid RAG application with cloud vector storage, API-based semantic embeddings, metadata-preserving parsing, semantic chunking, automatic source selection and inspectable citations. The remaining design boundary is visual OCR and mathematical image extraction on a low-memory hosted plan; the rest of the core document-text, table and retrieval flow is designed to run reliably without the developer laptop's GPU.")
    add_text(doc, "This report reflects the latest implementation files and the development changes made to improve accuracy, usability and deployment stability.", 10, MUTED, italic=True, before=10)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Document Q&A Chatbot - Updated Project Report"
    doc.core_properties.subject = "Architecture, changes, workflow and validation"
    doc.core_properties.author = "Document Q&A Chatbot Project"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
