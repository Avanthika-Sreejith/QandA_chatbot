from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "Document_QA_Chatbot_Project_Report.docx"
FLOW = ROOT / "work" / "document_qa_process_flow.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "5B6573"
GREEN = "E9F5EE"
AMBER = "FFF4D6"
RED = "FCE8E6"
CONTENT_DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def font(run, size=11, color="000000", bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def add_para(doc, text="", style=None, before=0, after=6, align=None, color="000000", size=11, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        font(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    font(p.add_run(text), 11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    font(p.add_run(text), 11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), {1: 16, 2: 13, 3: 12}[level], {1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return p


def add_table(doc, headers, rows, widths, status_column=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, value in zip(header.cells, headers):
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(value), 10, NAVY, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, (cell, value) in enumerate(zip(cells, row_values)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            text_color = "000000"
            fill = None
            if status_column is not None and idx == status_column:
                if "Pass" in value:
                    fill, text_color = GREEN, "1B5E3C"
                elif "Expected" in value or "Manual" in value:
                    fill, text_color = AMBER, "735300"
                elif "Limitation" in value:
                    fill, text_color = RED, "8A1C1C"
            if fill:
                set_cell_shading(cell, fill)
            font(p.add_run(value), 9.4, text_color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    font(run, 9, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def set_image_alt(inline_shape, description):
    """Add an accessible description to an inline image."""
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", "Document Q&A processing flow")


def build_flow_image():
    img = Image.new("RGB", (1600, 1030), "#FFFFFF")
    d = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 42)
        box_font = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 31)
        small_font = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 24)
    except OSError:
        title_font = box_font = small_font = ImageFont.load_default()
    d.text((800, 35), "Document Q&A Chatbot: End-to-End Process Flow", anchor="ma", font=title_font, fill="#0B2545")
    boxes = [
        ("1. User input", "Upload files or a ZIP folder", "#E8EEF5"),
        ("2. Parse and chunk", "PDF/DOCX/XLSX/XLSM/TXT -> text + metadata -> 800-character chunks", "#E9F5EE"),
        ("3. Create vectors", "Jina dense embedding API + FastEmbed BM25 sparse embeddings", "#FFF4D6"),
        ("4. Store and retrieve", "Qdrant Cloud stores vectors; search is filtered to the active document chat", "#E8EEF5"),
        ("5. Answer and persist", "Groq Llama 3.3 70B generates a grounded answer; Supabase saves Q&A history", "#E9F5EE"),
    ]
    y = 130
    for title, detail, fill in boxes:
        x1, x2 = 150, 1450
        d.rounded_rectangle((x1, y, x2, y + 120), radius=20, fill=fill, outline="#2E74B5", width=3)
        d.text((200, y + 28), title, font=box_font, fill="#0B2545")
        d.text((200, y + 73), detail, font=small_font, fill="#263238")
        if y < 130 + 4 * 170:
            d.line((800, y + 122, 800, y + 160), fill="#2E74B5", width=5)
            d.polygon([(786, y + 151), (814, y + 151), (800, y + 169)], fill="#2E74B5")
        y += 170
    img.save(FLOW)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in [(1, 16, BLUE, 16, 8), (2, 13, BLUE, 12, 6), (3, 12, DARK_BLUE, 8, 4)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("DOCUMENT Q&A CHATBOT | PROJECT DOCUMENTATION"), 9, MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    return doc


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build_flow_image()
    doc = setup_document()

    # Editorial-cover title block.
    add_para(doc, "PROJECT DOCUMENTATION", before=90, after=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, size=12, bold=True)
    add_para(doc, "Document Q&A Chatbot", after=7, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY, size=30, bold=True)
    add_para(doc, "A Retrieval-Augmented Generation System for Document-Based Question Answering", after=34, align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, size=14)
    cover = doc.add_table(rows=3, cols=2)
    cover.style = "Table Grid"
    set_table_geometry(cover, [2700, 6660])
    set_repeat_table_header(cover.rows[0])
    for cell, label, value in [
        (cover.cell(0, 0), "Prepared by", "Avanthika Sreejith"),
        (cover.cell(0, 1), "Project type", "RAG-based document intelligence application"),
        (cover.cell(1, 0), "Deployment", "Render (Docker web service)"),
        (cover.cell(1, 1), "Vector database", "Qdrant Cloud"),
        (cover.cell(2, 0), "Persistence", "Supabase"),
        (cover.cell(2, 1), "Report date", date.today().strftime("%d %B %Y")),
    ]:
        cell.text = ""
        p = cell.paragraphs[0]
        font(p.add_run(label + "\n"), 9, MUTED, bold=True)
        font(p.add_run(value), 10.5, NAVY, bold=True)
        set_cell_shading(cell, LIGHT_BLUE)
    add_para(doc, "This report documents the system design, processing flow, implementation choices, deployment architecture, and edge-case testing performed for the Document Q&A Chatbot.", before=34, after=0, align=WD_ALIGN_PARAGRAPH.CENTER, color="303840", size=11)
    doc.add_page_break()

    add_heading(doc, "1. Executive Summary")
    add_para(doc, "The Document Q&A Chatbot is a web application that lets users upload documents and ask natural-language questions about their contents. It follows a Retrieval-Augmented Generation (RAG) architecture: source text is extracted, divided into retrieval-sized chunks, converted to vector representations, stored in Qdrant Cloud, retrieved for a question, and supplied to a large language model to generate a grounded answer with source references.")
    add_para(doc, "The product is organized around document chats. Each chat has its own indexed documents and question history. This separation prevents a question in one chat from retrieving content indexed under another chat, while retaining earlier document collections for later use.")

    add_heading(doc, "2. Problem Statement and Objectives")
    add_para(doc, "Users often need answers from long, mixed-format files without reading every page or manually locating the relevant passage. Generic chatbots may answer from general knowledge rather than the supplied files. The project addresses this by retrieving source passages before generating every answer.")
    for item in [
        "Accept common business and academic document formats: PDF, DOCX, XLSX, XLSM, and TXT.",
        "Restrict retrieval to the currently selected document chat.",
        "Generate concise answers based only on retrieved passages and show source filenames.",
        "Persist document-chat names and Q&A history across restarts.",
        "Support deployment without local Qdrant Docker by using Qdrant Cloud.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. System Architecture")
    add_para(doc, "The architecture is intentionally modular. Streamlit presents the user interface, parsers convert files to text, Chonkie prepares chunks, embedding components create dense and sparse vectors, Qdrant Cloud manages vector storage, Groq produces the final answer, and Supabase stores the document-chat and message history.")
    flow_shape = doc.add_picture(str(FLOW), width=Inches(6.5))
    set_image_alt(
        flow_shape,
        "Flow diagram showing document input, parsing and chunking, dense and sparse embedding generation, Qdrant retrieval filtered to the active chat, and Groq answer generation with Supabase persistence.",
    )
    caption = add_para(doc, "Figure 1. End-to-end document ingestion, retrieval, and answer-generation flow.", before=3, after=8, align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, size=9, italic=True)

    add_heading(doc, "4. Technology Stack")
    add_table(doc,
        ["Layer", "Technology", "Role in the system"],
        [
            ("User interface", "Streamlit", "Provides document-chat management, uploads, progress feedback, and the chat interface."),
            ("Document parsing", "PyMuPDF, python-docx, openpyxl", "Extracts textual content and source metadata from PDFs, Word files, spreadsheets, and TXT files."),
            ("Chunking", "Chonkie RecursiveChunker", "Splits extracted text into 800-character chunks at natural boundaries where possible."),
            ("Dense embeddings", "Jina Embeddings v3 API", "Creates 1,024-dimensional semantic vectors using an OpenAI-compatible API."),
            ("Sparse embeddings", "FastEmbed - Qdrant/bm25", "Creates BM25-style sparse vectors stored alongside dense vectors."),
            ("Vector store", "Qdrant Cloud", "Stores chunks, vectors, metadata, and chat-scoping payloads."),
            ("Answer generation", "Groq - Llama 3.3 70B Versatile", "Generates a grounded response from the retrieved context passages."),
            ("Persistence", "Supabase", "Stores document-chat metadata and question/answer history."),
            ("Hosting", "Render + Docker", "Runs the Streamlit application as a hosted web service."),
        ], [1900, 2450, 5010])

    add_heading(doc, "5. Detailed Processing Flow")
    add_heading(doc, "5.1 Ingestion Flow", 2)
    for item in [
        "The user selects one or more supported files or uploads a ZIP archive containing a folder and its subfolders.",
        "The application stores uploaded files in a working upload directory and invokes the relevant parser. Each parser returns text-bearing segments plus metadata such as filename, path, page, sheet, paragraph, or section.",
        "Chonkie recursively splits the segments into chunks of up to 800 characters, preferring paragraph, sentence, and word boundaries before a character-level split.",
        "For every batch of chunks, the Jina embedding API creates dense 1,024-dimensional vectors. FastEmbed creates sparse BM25 vectors for the same chunks.",
        "The application upserts each chunk and both vector forms into Qdrant Cloud. Payload metadata includes the document-chat identifier and chat name.",
        "After indexing, the application refreshes the indexed-file list and clears the upload selection so users can add more documents without confusion.",
    ]:
        add_number(doc, item)

    add_heading(doc, "5.2 Query and Answer Flow", 2)
    for item in [
        "The user enters a question in the chat input. Blank input is not submitted.",
        "The Jina embedding API converts the question into a dense query vector.",
        "Qdrant Cloud performs semantic vector search and applies a payload filter for document_chat_id, ensuring that only the active chat's documents are eligible for retrieval.",
        "The top matching, de-duplicated passages are passed to Groq together with a system instruction that requires answers to be based only on the supplied context.",
        "The generated answer and source list are displayed in the conversation and persisted to Supabase for the selected chat.",
    ]:
        add_number(doc, item)
    add_para(doc, "Implementation note: dense and sparse vectors are both created and stored in Qdrant. The current query path uses dense semantic retrieval with chat filtering. Score fusion/hybrid query execution is a planned enhancement, rather than a capability claimed by the current search function.", before=6, after=8, color="735300", italic=True)

    add_heading(doc, "6. Document Chats, Isolation, and Persistence")
    add_para(doc, "A document chat is a persistent knowledge space. It has a unique identifier, a user-editable name, an indexed-file collection, and its own stored Q&A messages. The UI supports creating a new chat, automatically naming an initially untitled chat after its first upload, manually renaming it, clearing its question history, and permanently deleting a chat.")
    add_table(doc,
        ["Requirement", "Implementation"],
        [
            ("Query isolation", "Every Qdrant point contains document_chat_id. Retrieval filters on that payload field."),
            ("Chat history", "Supabase table chat_messages stores question, answer, source metadata, and timestamped order."),
            ("Chat metadata", "Supabase table document_chats stores the chat identifier, name, and creation details."),
            ("Deletion", "The application deletes Qdrant points belonging to the chat and deletes the Supabase chat row; related messages are removed by database cascade."),
            ("Indexed files", "The UI scrolls Qdrant payloads for the active chat and displays unique file paths in an expandable list."),
        ], [2500, 6860])

    add_heading(doc, "7. Deployment Architecture")
    add_para(doc, "The application is packaged with Docker and deployed as a Render web service. Render runs Streamlit and exposes the service port. Qdrant Cloud, Supabase, Jina, and Groq are external managed services accessed through API credentials stored as environment variables rather than hard-coded values.")
    add_table(doc,
        ["Environment variable", "Purpose"],
        [
            ("QDRANT_URL / QDRANT_API_KEY", "Connect to the Qdrant Cloud vector collection."),
            ("QDRANT_COLLECTION", "Select the collection used for document chunks."),
            ("JINA_API_KEY", "Authorize dense embedding requests to the Jina API."),
            ("GROQ_API_KEY", "Authorize answer-generation requests to Groq."),
            ("SUPABASE_URL / SUPABASE_KEY", "Connect to Supabase for chat and message persistence."),
        ], [3450, 5910])
    add_para(doc, "Because the hosted Render service cannot read arbitrary paths such as D:\\test on a user's computer, hosted users use multi-file upload or the ZIP-folder upload option. The local-folder-path option remains useful only when Streamlit is running on the same machine as the folder.", before=6, after=8, color="735300", italic=True)

    add_heading(doc, "8. Edge-Case Testing")
    add_para(doc, "The following manual test cases were included to validate input handling, fault tolerance, document-chat isolation, and retrieval behavior. The matrix records the expected behavior and the relevant application handling; it should be retained with the project submission as test evidence.")
    add_heading(doc, "8.1 File Upload and Parsing Tests", 2)
    add_table(doc,
        ["Test case", "Expected behavior / implementation response", "Status"],
        [
            ("Empty PDF, DOCX, or TXT", "No crash. The file produces no extractable segments and the UI reports that no text could be extracted.", "Manual test"),
            ("Scanned or image-only PDF", "No OCR is performed in V1. If no text layer exists, it is reported as having no extractable text.", "Manual test"),
            ("Large file (100 MB+)", "The browser/Streamlit upload limit and hosting resources determine acceptance. The UI must show a clear upload or indexing error rather than crash.", "Manual test"),
            ("Special characters in filename", "The uploaded filename is handled as a path component and parser metadata preserves the source name/path.", "Manual test"),
            ("Corrupted PDF, DOCX, or XLSX", "The indexing operation surfaces an error message in the UI; the Streamlit application remains available.", "Manual test"),
            ("PNG or JPG in normal file upload", "Rejected by the normal file-uploader type filter because image OCR is outside V1 scope.", "Pass"),
            ("ZIP archive", "Accepted only in the dedicated ZIP-folder uploader. Supported documents inside nested folders are extracted safely and indexed.", "Pass"),
            ("10+ files in one upload", "Files are parsed sequentially and embedded in batches, with progress feedback shown during processing.", "Manual test"),
            ("Same file uploaded twice in one chat", "Stable point identifiers cause matching chunks to be upserted (re-indexed) rather than creating duplicate points for the same path/chunk.", "Pass"),
            ("DOCX containing only images", "No text paragraphs or text tables are extracted, so the file returns zero segments and is not indexed.", "Pass"),
        ], [2600, 5000, 1760], status_column=2)

    add_heading(doc, "8.2 Search and Retrieval Tests", 2)
    add_table(doc,
        ["Test case", "Expected behavior / implementation response", "Status"],
        [
            ("Empty query", "The chat input does not submit a blank or whitespace-only question.", "Pass"),
            ("Question not covered by documents", "The answer prompt instructs the model to state that context is insufficient. A future relevance threshold can make this stricter.", "Manual test"),
            ("Special characters / SQL-like text", "The query is sent as embedding input and Qdrant uses typed filters; user text is not constructed as SQL.", "Pass"),
            ("Chat with no indexed files", "Qdrant returns no matching passages; the answer workflow responds gracefully with no relevant context.", "Pass"),
            ("Question across different chats", "The Qdrant document_chat_id filter restricts retrieval to the active document chat only.", "Pass"),
            ("Very long query (500+ words)", "The question is embedded and searched subject to API/provider input limits; the UI surfaces any provider error clearly.", "Manual test"),
        ], [2600, 5000, 1760], status_column=2)

    add_heading(doc, "9. Limitations and Future Enhancements")
    for item in [
        "OCR is not included in Version 1; scanned PDFs and image-only Word files have no extractable text.",
        "A browser cannot provide an arbitrary local folder path to the hosted application. ZIP folder upload provides a consistent folder-import workflow for every user.",
        "The current retrieval call is dense semantic search. Stored BM25 sparse vectors enable a future hybrid fusion/reranking implementation.",
        "A relevance-score threshold and explicit 'no answer' guard can further reduce answers from weakly related passages.",
        "Raw uploads stored on a hosted container may be temporary; Qdrant vectors and Supabase chat records are the durable application state.",
        "Future improvements can include OCR, true hybrid search, reranking, automated regression tests, per-file deletion, rate limiting, and user authentication.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10. Security and Data Handling")
    add_para(doc, "API keys are supplied through environment variables and should never be committed to source control. The Qdrant API key restricts vector-database access, while Supabase credentials control persistence. The implementation uses a chat identifier as a Qdrant payload filter to separate document collections at retrieval time. Production hardening should additionally include authentication, authorization checks, upload malware scanning, file-retention rules, and rate limiting.")

    add_heading(doc, "11. Conclusion")
    add_para(doc, "The Document Q&A Chatbot demonstrates an end-to-end RAG implementation that turns uploaded documents into a persistent, chat-scoped knowledge base. Its modular design separates parsing, chunking, embeddings, vector storage, retrieval, generation, and persistence. The deployment removes the need for users to run local Qdrant Docker, while the documented test matrix captures important edge cases for continued validation and improvement.")

    doc.core_properties.title = "Document Q&A Chatbot - Project Documentation"
    doc.core_properties.author = "Avanthika Sreejith"
    doc.core_properties.subject = "RAG-based document question-answering system"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
