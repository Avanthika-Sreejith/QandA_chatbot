"""Extract text and source metadata from supported document types."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any
import fitz
import math
import re
import tempfile
from zipfile import ZipFile

from app.config import ENABLE_PDF_OCR


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xlsm", ".txt"}
TABLE_ROWS_PER_SEGMENT = 20
MIN_OCR_IMAGE_SIZE = (25, 12)  # ignore tiny icons/decorative marks
STANDALONE_OCR_IMAGE_SIZE = (60, 12)  # ignore logos/borders; capture inline formula images
MAX_OCR_IMAGES_PER_TABLE = 6
VECTOR_GLYPH_MIN_SIZE = (4, 4)  # ignore bullets/decorative marks
VECTOR_GLYPH_MAX_SIZE = (40, 40)  # ignore table borders / large fills
VECTOR_CLUSTER_Y_TOL = 6.0  # treat glyphs on the same text line as one region
VECTOR_OCR_PADDING = 2.0  # avoid clipping glyphs at region edges (standalone only)
VECTOR_TABLE_OCR_PADDING = 0.0  # table grid lines corrupt OCR if included

_OCR_ENGINE: Any = None


def _get_ocr_engine() -> Any:
    """Return a process-wide OCR engine, creating it once."""
    global _OCR_ENGINE
    if _OCR_ENGINE is None:
        from rapidocr_onnxruntime import RapidOCR

        _OCR_ENGINE = RapidOCR()
    return _OCR_ENGINE


def _looks_like_numbered_heading(text: str) -> bool:
    """Recognize simple headings whose DOCX style was not set to Heading."""
    return bool(re.fullmatch(r"\d+(?:\.\d+)*\.?\s+[A-Z][^.!?]{0,100}", text))


def _heading_level(style_name: str) -> int:
    """Extract the heading level from a style name such as 'Heading 2'."""
    match = re.search(r"(\d+)", style_name)
    return int(match.group(1)) if match else 1


@dataclass(frozen=True)
class ParsedSegment:
    """A source-preserving unit of extracted text before chunking."""

    text: str
    metadata: dict[str, Any]


def _base_metadata(path: Path) -> dict[str, Any]:
    return {"file_name": path.name, "file_path": str(path.resolve())}


def _table_row_groups(
    rows: list[list[str]],
    *,
    path: Path,
    source_type: str,
    table_number: int | None = None,
    sheet: str | None = None,
    has_embedded_images: bool = False,
    image_ocr: list[str] | None = None,
    page: int | None = None,
) -> list[ParsedSegment]:
    """Create header-preserving chunks for a native document table."""
    non_empty = [row for row in rows if any(cell.strip() for cell in row)]
    if not non_empty:
        return []

    header_values = non_empty[0]
    headers = [value or f"Column {index}" for index, value in enumerate(header_values, start=1)]
    data_rows = non_empty[1:] or [non_empty[0]]
    ocr_texts = [text for text in (image_ocr or []) if text.strip()]
    output: list[ParsedSegment] = []
    for start in range(0, len(data_rows), TABLE_ROWS_PER_SEGMENT):
        group = data_rows[start : start + TABLE_ROWS_PER_SEGMENT]
        rendered_rows = []
        for row_offset, row in enumerate(group, start=start + 1):
            values = (row + [""] * len(headers))[: len(headers)]
            cells = [f"{header}: {value}" for header, value in zip(headers, values, strict=True) if value]
            if cells:
                rendered_rows.append(f"Row {row_offset}: " + "; ".join(cells))
        if not rendered_rows and not ocr_texts:
            continue

        label = f"Table {table_number}" if table_number is not None else f"Sheet {sheet}"
        text = f"{label}\nHeaders: " + " | ".join(headers) + "\n" + "\n".join(rendered_rows)
        if ocr_texts:
            text += "\n\n[Image content]\n" + "\n".join(ocr_texts)
        metadata = _base_metadata(path) | {
            "source_type": source_type,
            "content_kind": "table",
            "table_row_start": start + 1,
            "table_row_end": start + len(group),
            "table_headers": headers,
            "has_embedded_images": has_embedded_images,
            "image_ocr": ocr_texts,
        }
        if page is not None:
            metadata["page"] = page
        if table_number is not None:
            metadata["table"] = table_number
            metadata["section"] = label
        if sheet is not None:
            metadata["sheet"] = sheet
        output.append(ParsedSegment(text, metadata))
    return output


def _median_font_size(spans: list[dict[str, Any]]) -> float:
    """Return the median span size, used as the body text size."""
    sizes = sorted(span.get("size", 0.0) for span in spans)
    if not sizes:
        return 0.0
    return sizes[len(sizes) // 2]


def _looks_like_pdf_heading(text: str, max_size: float, is_bold: bool, body_size: float) -> bool:
    """Heuristically detect a PDF heading line from its size and weight."""
    if len(text) < 2 or len(text) > 120:
        return False
    if text.strip().isdigit():  # skip page numbers
        return False
    if body_size <= 0:
        return is_bold and len(text) <= 80
    larger_font = max_size >= body_size + 1.5
    # Many study-note PDFs use bold, body-sized Title Case labels such as
    # "Key concepts". Preserve them as headings even without a larger font so
    # their terms are included with the following semantic chunk.
    words = text.split()
    # PDF exports often capitalise only the first word of a subsection label
    # (for example, "Key concepts"), so str.title() is too strict here.
    # Short labels without trailing punctuation are safe to treat as headings;
    # definition lines such as "Normal time:" are excluded by their colon.
    title_case_label = (
        2 <= len(words) <= 8
        and text[0].isupper()
        and not any(char in text for char in ".,:;!?()")
    )
    return (larger_font and len(text) <= 80) or (is_bold and len(text) <= 60) or title_case_label


def _normalise_cells(row: list[str | None] | None) -> list[str]:
    """Clean a raw PDF table row (None cells become empty strings)."""
    if not row:
        return []
    return [" ".join((cell or "").split()) for cell in row]


def _line_overlaps_tables(line_rect: Any, table_boxes: list[Any]) -> bool:
    """Return True when a text line sits mostly inside a detected table region."""
    if line_rect is None or not table_boxes:
        return False
    line_area = line_rect.get_area()
    if line_area <= 0:
        return False
    for table_rect in table_boxes:
        intersection = line_rect & table_rect
        if not intersection.is_empty and intersection.get_area() / line_area > 0.5:
            return True
    return False


def _image_overlaps_tables(image_rect: Any, table_boxes: list[Any]) -> bool:
    """Return True when an image sits mostly inside a detected table region."""
    if not table_boxes:
        return False
    image_area = image_rect.get_area()
    if image_area <= 0:
        return False
    for table_rect in table_boxes:
        intersection = image_rect & table_rect
        if not intersection.is_empty and intersection.get_area() / image_area > 0.5:
            return True
    return False


def _text_line_rects(page: Any) -> list[Any]:
    """Bounding boxes of every text line on the page."""
    rects: list[Any] = []
    for block in page.get_text("dict").get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            bbox = line.get("bbox")
            if bbox:
                rects.append(fitz.Rect(bbox))
    return rects


def _vector_glyph_rects(
    page: Any, table_boxes: list[Any], inside_table: bool, text_lines: list[Any]
) -> list[Any]:
    """Vector-drawn glyph rects that resemble formula characters.

    PyMuPDF's ``get_image_info()`` only reports raster images, so formulas
    drawn as vector paths are invisible to the existing OCR path. This collects
    small, square-ish vector paths that are not covered by text (which would be
    underlines, highlights, or bullets) and filters by whether they sit inside a
    detected table region.
    """
    glyphs: list[Any] = []
    for drawing in page.get_drawings():
        rect = fitz.Rect(drawing.get("rect"))
        if not (VECTOR_GLYPH_MIN_SIZE[0] <= rect.width <= VECTOR_GLYPH_MAX_SIZE[0]):
            continue
        if not (VECTOR_GLYPH_MIN_SIZE[1] <= rect.height <= VECTOR_GLYPH_MAX_SIZE[1]):
            continue
        # Skip glyphs mostly covered by a text line (bullets, underlines, etc.).
        if rect.get_area() > 0:
            covered = False
            for line_rect in text_lines:
                intersection = rect & line_rect
                if not intersection.is_empty and intersection.get_area() / rect.get_area() > 0.6:
                    covered = True
                    break
            if covered:
                continue
        is_in_table = _image_overlaps_tables(rect, table_boxes)
        if is_in_table != inside_table:
            continue
        glyphs.append(rect)
    return glyphs


def _cluster_vector_glyphs(glyphs: list[Any]) -> list[Any]:
    """Merge glyphs sharing a text line into one OCR region."""
    clusters: list[Any] = []
    for rect in sorted(glyphs, key=lambda r: (r.y0, r.x0)):
        merged = False
        for cluster in clusters:
            if abs(rect.y0 - cluster.y0) <= VECTOR_CLUSTER_Y_TOL or abs(
                rect.y1 - cluster.y1
            ) <= VECTOR_CLUSTER_Y_TOL:
                cluster.x0 = min(cluster.x0, rect.x0)
                cluster.y0 = min(cluster.y0, rect.y0)
                cluster.x1 = max(cluster.x1, rect.x1)
                cluster.y1 = max(cluster.y1, rect.y1)
                merged = True
                break
        if not merged:
            clusters.append(fitz.Rect(rect))
    return clusters


def _is_section_name_line(text: str) -> bool:
    """True for short, punctuation-free, Title-Case lines (section titles)."""
    if not (2 <= len(text) <= 45):
        return False
    if any(char in text for char in ".,:;!?()"):
        return False
    if not text[0].isupper():
        return False
    if len(text.split()) < 2:
        return False
    return bool(re.fullmatch(r"[A-Za-z][A-Za-z &'\-/]*", text))


def _drop_section_name_runs(lines: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Remove runs of 3+ consecutive bare section-title lines (navigation lists)."""
    flags = [
        _is_section_name_line(
            "".join(span.get("text", "") for span in line.get("spans", [])).strip()
        )
        for line in lines
    ]
    keep = [True] * len(lines)
    i = 0
    while i < len(flags):
        if flags[i]:
            j = i
            while j < len(flags) and flags[j]:
                j += 1
            if j - i >= 3:
                for k in range(i, j):
                    keep[k] = False
            i = j
        else:
            i += 1
    return [line for line, keep_line in zip(lines, keep) if keep_line]


def _ocr_image(
    page: Any,
    image_rect: Any,
    ocr_engine: Any,
    min_size: tuple[int, int],
    min_zoom: float = 0.0,
) -> str | None:
    """OCR one rendered image region and return its recognised text.

    Renders the region at a zoom high enough for small inline equations and
    returns None when the image is too small, OCR fails, or the recognised text
    is too short to be meaningful (icons/logos).
    """
    # RapidOCR/ONNX consumes more memory than a 512 MB Render instance can
    # safely provide. Keep the main RAG path reliable unless OCR was
    # deliberately enabled on a larger deployment.
    if not ENABLE_PDF_OCR:
        return None
    if image_rect.width < min_size[0] or image_rect.height < min_size[1]:
        return None
    try:
        ocr_engine = _get_ocr_engine()
        # Target a rendered text height of ~48px so tiny equations stay legible.
        zoom = max(min_zoom, min(8, max(3, math.ceil(48 / max(image_rect.height, 1)))))
        pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), clip=image_rect)
        # A named temp file keeps its handle open on Windows, which blocks the
        # writer and the OCR engine; write into a fresh temp directory instead.
        tmp_dir = Path(tempfile.mkdtemp(prefix="rag_ocr_"))
        tmp_path = tmp_dir / "image.png"
        pixmap.save(str(tmp_path))
        try:
            result, _ = ocr_engine(str(tmp_path))
        finally:
            tmp_path.unlink(missing_ok=True)
            tmp_dir.rmdir()
        if result:
            text = " ".join(item[1] for item in result).strip()
            if len(text) >= 6 and len(text.split()) >= 2:
                return text
    except Exception as error:
        import logging

        logging.getLogger("ocr").warning(
            "OCR failed for region %s: %s", tuple(round(v, 1) for v in image_rect), error
        )
        return None
    return None


def parse_pdf(path: Path) -> list[ParsedSegment]:
    """Return per-page segments: headings, prose sections, and detected tables."""
    import fitz

    segments: list[ParsedSegment] = []
    with fitz.open(path) as document:
        for index, page in enumerate(document, start=1):
            blocks = page.get_text("dict").get("blocks", [])
            all_lines = [
                line
                for block in blocks
                if block.get("type") == 0
                for line in block.get("lines", [])
            ]
            image_infos = page.get_image_info()
            if not all_lines and not image_infos:
                continue
            text_line_rects = _text_line_rects(page)

            # Detect native tables first; index them as header-preserving row groups.
            table_boxes: list[fitz.Rect] = []
            table_segments: list[ParsedSegment] = []
            ocr_engine: Any = None
            try:
                table_finder = page.find_tables()
            except Exception:
                table_finder = None
            if table_finder is not None:
                for table_number, table in enumerate(table_finder.tables, start=1):
                    table_rect = fitz.Rect(table.bbox)
                    rows = [_normalise_cells(row) for row in (table.extract() or [])]
                    intersecting = [
                        info
                        for info in image_infos
                        if fitz.Rect(info["bbox"]).intersects(table_rect)
                    ]
                    ocr_texts: list[str] = []
                    if intersecting:
                        has_rows = any(cell.strip() for row in rows for cell in row)
                        table_area = table_rect.get_area()
                        for info in intersecting[:MAX_OCR_IMAGES_PER_TABLE]:
                            image_rect = fitz.Rect(info["bbox"])
                            # Skip a screenshot that just duplicates the whole table.
                            if has_rows and table_area > 0 and image_rect.get_area() / table_area > 0.5:
                                continue
                            ocr_text = _ocr_image(page, image_rect, ocr_engine, MIN_OCR_IMAGE_SIZE)
                            if ocr_text:
                                ocr_texts.append(ocr_text)
                    # Formulas rendered as vector paths are invisible to
                    # get_image_info(); OCR glyph clusters inside the table.
                    table_glyphs = _vector_glyph_rects(
                        page, [table_rect], inside_table=True, text_lines=text_line_rects
                    )
                    for cluster in _cluster_vector_glyphs(table_glyphs)[
                        :MAX_OCR_IMAGES_PER_TABLE
                    ]:
                        padded = fitz.Rect(
                            cluster.x0 - VECTOR_TABLE_OCR_PADDING,
                            cluster.y0 - VECTOR_TABLE_OCR_PADDING,
                            cluster.x1 + VECTOR_TABLE_OCR_PADDING,
                            cluster.y1 + VECTOR_TABLE_OCR_PADDING,
                        )
                        cluster_zoom = max(
                            3, math.ceil(48 / max(cluster.height, 1))
                        )
                        ocr_text = _ocr_image(
                            page,
                            padded,
                            ocr_engine,
                            VECTOR_GLYPH_MIN_SIZE,
                            min_zoom=cluster_zoom,
                        )
                        if ocr_text:
                            ocr_texts.append(ocr_text)
                    candidate_segments = _table_row_groups(
                        rows,
                        path=path,
                        source_type="pdf_table",
                        table_number=table_number,
                        has_embedded_images=bool(intersecting or table_glyphs),
                        image_ocr=ocr_texts,
                        page=index,
                    )
                    # A visual divider/list can be misidentified as a PDF
                    # table. Only remove its text from the prose path when a
                    # real, non-empty table chunk was produced.
                    if candidate_segments:
                        table_boxes.append(table_rect)
                        table_segments.extend(candidate_segments)

            # OCR images that are not covered by any detected native table, such
            # as screenshot tables that find_tables cannot detect.
            standalone_count = 0
            for info in image_infos:
                if standalone_count >= MAX_OCR_IMAGES_PER_TABLE:
                    break
                image_rect = fitz.Rect(info["bbox"])
                if _image_overlaps_tables(image_rect, table_boxes):
                    continue
                ocr_text = _ocr_image(page, image_rect, ocr_engine, STANDALONE_OCR_IMAGE_SIZE)
                if not ocr_text:
                    continue
                standalone_count += 1
                segments.append(
                    ParsedSegment(
                        f"[Image content]\n{ocr_text}",
                        _base_metadata(path)
                        | {
                            "source_type": "pdf_image",
                            "content_kind": "image",
                            "page": index,
                            "section": "Image content",
                        },
                    )
                )

            # Also OCR formulas drawn as vector paths (e.g. MathType/MathML
            # output) that get_image_info() cannot see.
            standalone_glyphs = _vector_glyph_rects(
                page, table_boxes, inside_table=False, text_lines=text_line_rects
            )
            for cluster in _cluster_vector_glyphs(standalone_glyphs):
                if standalone_count >= MAX_OCR_IMAGES_PER_TABLE:
                    break
                padded = fitz.Rect(
                    cluster.x0 - VECTOR_OCR_PADDING,
                    cluster.y0 - VECTOR_OCR_PADDING,
                    cluster.x1 + VECTOR_OCR_PADDING,
                    cluster.y1 + VECTOR_OCR_PADDING,
                )
                cluster_zoom = max(3, math.ceil(48 / max(cluster.height, 1)))
                ocr_text = _ocr_image(
                    page,
                    padded,
                    ocr_engine,
                    VECTOR_GLYPH_MIN_SIZE,
                    min_zoom=cluster_zoom,
                )
                if not ocr_text:
                    continue
                standalone_count += 1
                segments.append(
                    ParsedSegment(
                        f"[Image content]\n{ocr_text}",
                        _base_metadata(path)
                        | {
                            "source_type": "pdf_image",
                            "content_kind": "image",
                            "page": index,
                            "section": "Image content",
                        },
                    )
                )

            # Body font size is measured from text outside the detected tables.
            non_table_lines = [
                line
                for line in all_lines
                if not _line_overlaps_tables(fitz.Rect(line.get("bbox")), table_boxes)
            ]
            non_table_lines = _drop_section_name_runs(non_table_lines)
            non_table_spans = [
                span
                for line in non_table_lines
                for span in line.get("spans", [])
            ]
            if not non_table_spans and not table_segments:
                continue
            all_spans = [
                span
                for line in all_lines
                for span in line.get("spans", [])
            ]
            body_size = (
                _median_font_size(non_table_spans)
                if non_table_spans
                else _median_font_size(all_spans)
            )

            current_section = "Document body"
            pending: list[str] = []
            page_segments: list[ParsedSegment] = []

            def flush() -> None:
                text = " ".join(pending).strip()
                if text:
                    # Add the parent heading to the embed text. This lets a
                    # query such as "key concepts in crashing" find the
                    # definitions under that heading rather than only a later
                    # broad summary of the topic.
                    if current_section != "Document body":
                        text = f"{current_section}\n{text}"
                    page_segments.append(
                        ParsedSegment(
                            text,
                            _base_metadata(path)
                            | {"source_type": "pdf", "page": index, "section": current_section},
                        )
                    )
                pending.clear()

            for line in non_table_lines:
                line_spans = line.get("spans", [])
                if not line_spans:
                    continue
                text = "".join(span.get("text", "") for span in line_spans).strip()
                if not text:
                    continue
                max_size = max((span.get("size", 0.0) for span in line_spans), default=0.0)
                is_bold = any(span.get("flags", 0) & 16 for span in line_spans)
                if _looks_like_pdf_heading(text, max_size, is_bold, body_size):
                    flush()
                    current_section = text
                    page_segments.append(
                        ParsedSegment(
                            text,
                            _base_metadata(path)
                            | {
                                "source_type": "pdf_heading",
                                "content_kind": "heading",
                                "page": index,
                                "section": text,
                                "is_heading": True,
                            },
                        )
                    )
                else:
                    pending.append(text)
            flush()
            segments.extend(table_segments)
            segments.extend(page_segments)
    return segments


def parse_docx(path: Path) -> list[ParsedSegment]:
    """Return paragraphs and tables, labelled by their nearest heading where possible."""
    from docx import Document

    document = Document(path)
    section = "Document body"
    segments: list[ParsedSegment] = []
    for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        is_heading = style_name.lower().startswith("heading") or _looks_like_numbered_heading(text)
        if is_heading:
            section = text
            segments.append(
                ParsedSegment(
                    text,
                    _base_metadata(path)
                    | {
                        "source_type": "docx_heading",
                        "content_kind": "heading",
                        "section": text,
                        "paragraph": paragraph_number,
                        "heading_level": _heading_level(style_name),
                        "is_heading": True,
                    },
                )
            )
            continue
        segments.append(
            ParsedSegment(
                text,
                _base_metadata(path)
                | {"source_type": "docx", "section": section, "paragraph": paragraph_number},
            )
        )

    # Preserve table headers in each row group. This is more useful for RAG
    # than flattening an entire table into one unstructured block of text.
    for table_number, table in enumerate(document.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(cells)
        image_count = len(table._element.xpath(".//a:blip"))
        segments.extend(
            _table_row_groups(
                rows,
                path=path,
                source_type="docx_table",
                table_number=table_number,
                has_embedded_images=bool(image_count),
            )
        )
    return segments


def parse_excel(path: Path) -> list[ParsedSegment]:
    """Return header-aware row groups and retain formulas plus cached values."""
    from openpyxl import load_workbook

    formula_workbook = load_workbook(path, read_only=True, data_only=False)
    value_workbook = load_workbook(path, read_only=True, data_only=True)
    with ZipFile(path) as archive:
        workbook_image_count = sum(name.startswith("xl/media/") for name in archive.namelist())
    segments: list[ParsedSegment] = []
    try:
        for worksheet, value_worksheet in zip(
            formula_workbook.worksheets, value_workbook.worksheets, strict=True
        ):
            rows: list[list[str]] = []
            for formula_row, value_row in zip(
                worksheet.iter_rows(values_only=True),
                value_worksheet.iter_rows(values_only=True),
                strict=True,
            ):
                values: list[str] = []
                for formula, calculated_value in zip(formula_row, value_row, strict=True):
                    if formula is None and calculated_value is None:
                        values.append("")
                    elif isinstance(formula, str) and formula.startswith("="):
                        values.append(f"Formula: {formula}; calculated value: {calculated_value!s}")
                    else:
                        values.append(str(calculated_value if calculated_value is not None else formula).strip())
                if any(values):
                    rows.append(values)
            segments.extend(
                _table_row_groups(
                    rows,
                    path=path,
                    source_type="excel_table",
                    sheet=worksheet.title,
                    has_embedded_images=bool(workbook_image_count),
                )
            )
    finally:
        formula_workbook.close()
        value_workbook.close()
    return segments


def parse_text(path: Path) -> list[ParsedSegment]:
    """Return a plain-text file as one source-preserving segment."""
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        return []
    return [
        ParsedSegment(
            text,
            _base_metadata(path)
            | {"source_type": "txt", "section": "Full document", "line_start": 1, "line_end": len(text.splitlines())},
        )
    ]


def parse_document(file_path: str | Path) -> list[ParsedSegment]:
    """Dispatch parsing based on the file extension."""
    path = Path(file_path)
    if not path.is_file():
        raise FileNotFoundError(f"File not found: {path}")
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type '{extension}'. Supported types: {supported}")
    if extension == ".pdf":
        return parse_pdf(path)
    if extension == ".docx":
        return parse_docx(path)
    if extension in {".xlsx", ".xlsm"}:
        return parse_excel(path)
    return parse_text(path)
